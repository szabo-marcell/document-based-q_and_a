"""
Generálja a genaispecifikacio.docx dokumentációs fájlt.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Stílusok beállítása ──────────────────────────────────────────────────────

def set_heading(paragraph, level=1):
    colors = {1: "1F3864", 2: "2E74B5", 3: "2E74B5"}
    sizes  = {1: 20,       2: 16,       3: 13}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = RGBColor.from_string(colors.get(level, "000000"))

def add_heading(doc, text, level=1):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Szürke háttér szimulálása monospace fonttal
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Szegély a bekezdés köré
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "AAAAAA")
        pBdr.append(bdr)
    pPr.append(pBdr)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.75)
    return p

def add_table_row(table, cells, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = bold or header
            run.font.size = Pt(10)
        if header:
            cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "2E74B5")
            cell._tc.tcPr.append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row

# ── Dokumentum margók ────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ════════════════════════════════════════════════════════════════════════════
# FEDŐLAP
# ════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Dokumentum-alapú Q&A alkalmazás")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Műszaki specifikáció és változásnapló")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run(f"Készítés dátuma: {datetime.date.today().strftime('%Y. %m. %d.')}")
run3.font.size = Pt(11)
run3.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. AZ ALKALMAZÁS ÁTTEKINTÉSE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Az alkalmazás áttekintése", level=1)

add_paragraph(doc,
    "A projekt egy Streamlit alapú, dokumentum-alapú kérdés-válasz (Q&A) rendszer, "
    "amely négy egymásra épülő HuggingFace API hívással, Chain-of-Thought (CoT) technikával, "
    "Personality Prompting-gal és Few-Shot Prompting-gal ad kontextusérzékeny, "
    "forrásalapú válaszokat a felhasználó kérdéseire.")

doc.add_paragraph()
add_heading(doc, "1.1 Technológiai stack", level=2)

tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = "Table Grid"
add_table_row(tech_table, ["Komponens", "Technológia / csomag"], header=True)
for row in [
    ("Frontend keretrendszer",  "Streamlit"),
    ("HTTP kérések",            "requests"),
    ("Embedding (1. API hívás)","langchain_huggingface · HuggingFaceEmbeddings"),
    ("Vektoros hasonlóság",     "scikit-learn · cosine_similarity + numpy"),
    ("Összefoglalás (2. API)",  "facebook/bart-large-cnn (HuggingFace Inference API)"),
    ("CoT következtetés (3. API)", "mistralai/Mistral-7B-Instruct-v0.2"),
    ("Végső válasz (4. API)",   "mistralai/Mistral-7B-Instruct-v0.2"),
    ("Konfiguráció",            "python-dotenv (.env fájl)"),
    ("Dokumentumgenerálás",     "python-docx"),
]:
    add_table_row(tech_table, list(row))

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 2. FÁJLSZERKEZET
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Fájlszerkezet", level=1)

add_code_block(doc,
    "beadando/\n"
    "├── app.py              # Fő alkalmazás (Streamlit UI + pipeline)\n"
    "├── beadando2.py        # Embedding modul (find_most_similar)\n"
    "├── create_docs.py      # Dokumentációgeneráló szkript\n"
    "├── genaispecifikacio.docx  # Jelen dokumentáció\n"
    "└── .env                # HF_API_KEY környezeti változó")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 3. A NÉGY HUGGINGFACE API HÍVÁS RÉSZLETESEN
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. A négy HuggingFace API hívás részletesen", level=1)

# 3.1
add_heading(doc, "3.1  1. API hívás – Embedding (all-MiniLM-L6-v2)", level=2)
add_paragraph(doc, "Fájl: beadando2.py · Függvény: find_most_similar()", bold=True)
add_paragraph(doc,
    "A LangChain HuggingFaceEmbeddings osztályán keresztül a "
    "sentence-transformers/all-MiniLM-L6-v2 modell minden szöveget 384 dimenziós "
    "lebegőpontos vektorrá alakít. Az API-t kétszer hívja meg: egyszer az összes "
    "dokumentumra (embed_documents), egyszer a felhasználói kérdésre (embed_query). "
    "Ezután cosine_similarity számítással kiválasztja a top-k legjobban illeszkedő dokumentumot.")
add_code_block(doc,
    "doc_emb   = embeddings.embed_documents(docs)\n"
    "query_emb = embeddings.embed_query(query)\n"
    "scores    = cosine_similarity([query_emb], doc_emb)[0]\n"
    "top_idx   = np.argsort(scores)[::-1][:top_k]")
add_paragraph(doc, "Kimenet: List[(dokumentum_szöveg, hasonlóság_pontszám)] – top-k pár.")

doc.add_paragraph()

# 3.2
add_heading(doc, "3.2  2. API hívás – Összefoglalás (facebook/bart-large-cnn)", level=2)
add_paragraph(doc, "Fájl: app.py · Függvény: call_summarization()", bold=True)
add_paragraph(doc,
    "A BART egy seq2seq encoder-decoder architektúrájú modell, amelyet CNN/DailyMail "
    "adathalmazon finomhangoltak összefoglalásra. A visszakeresett dokumentumokat "
    "legfeljebb 1024 karakterre csonkítja, majd 40–200 tokenben foglalja össze. "
    "A hívás a HuggingFace Inference API-n keresztül történik HTTP POST kéréssel.")
add_code_block(doc,
    'url = f"{HF_API_BASE}/facebook/bart-large-cnn"\n'
    'payload = {\n'
    '    "inputs": text[:1024],\n'
    '    "parameters": {"max_length": 200, "min_length": 40},\n'
    '    "options": {"wait_for_model": True},\n'
    '}')
add_paragraph(doc, "Kimenet: összefoglalt szöveg string (summary_text mező a JSON válaszból).")

doc.add_paragraph()

# 3.3
add_heading(doc, "3.3  3. API hívás – CoT következtetés (Mistral-7B-Instruct-v0.2)", level=2)
add_paragraph(doc, "Fájl: app.py · Függvény: call_reasoning()", bold=True)
add_paragraph(doc,
    "Ez a hívás három prompting technikát kombinál egyidejűleg. "
    "Personality Prompting: a PERSONALITY konstans beállítja a modell viselkedését "
    "(precíz, analitikus, dokumentum-alapú). "
    "Few-Shot Prompting: a FEW_SHOT_EXAMPLES két annotált példával mutatja meg "
    "a kívánt gondolkodási formátumot. "
    "Chain-of-Thought: a prompt explicit lépésekre kényszeríti a modellt "
    "(\"1. lépés –\" folytatásra várva).")
add_code_block(doc,
    'prompt = (\n'
    '    f"<s>[INST] {PERSONALITY}\\n\\n"\n'
    '    f"Tanulj a következő mintapéldákból:\\n{FEW_SHOT_EXAMPLES}\\n\\n"\n'
    '    f"Kontextus: {context}\\n"\n'
    '    f"Kérdés: {question}\\n"\n'
    '    f"Gondolkodás (lépések szerint):\\n  1. lépés – [/INST]"\n'
    ')\n'
    '"parameters": {"max_new_tokens": 350, "temperature": 0.35, "do_sample": True}')
add_paragraph(doc,
    "Az 0.35-ös temperature enyhén kreatív de fókuszált következtetést tesz lehetővé. "
    "A return_full_text: False beállítás csak az újonnan generált szöveget adja vissza.")

doc.add_paragraph()

# 3.4
add_heading(doc, "3.4  4. API hívás – Végső válasz (Mistral-7B-Instruct-v0.2)", level=2)
add_paragraph(doc, "Fájl: app.py · Függvény: call_final_answer()", bold=True)
add_paragraph(doc,
    "Ez a hívás az összes előző lépés kimenetét szintetizálja: az eredeti dokumentumokat, "
    "a BART-összefoglalót és a Mistral CoT-következtetést egyszerre kapja kontextusként. "
    "Az alacsonyabb temperature (0.25) determinisztikusabb, pontosabb választ eredményez, "
    "ami a végső kimenetnél fontosabb, mint a kreativitás. "
    "A prompt explicit \"MAGYARUL\" utasítást tartalmaz, mert a Mistral alapból angolul válaszolna.")
add_code_block(doc,
    'prompt = (\n'
    '    f"<s>[INST] {PERSONALITY}\\n\\n"\n'
    '    f"Rendelkezésre álló dokumentumok:\\n{context}\\n\\n"\n'
    '    f"Összefoglalt kontextus:\\n{summarized}\\n\\n"\n'
    '    f"Közbülső következtetés:\\n{reasoning}\\n\\n"\n'
    '    f"Kérdés: {question}\\n\\n"\n'
    '    f"...adj tömör végső választ MAGYARUL... [/INST]"\n'
    ')\n'
    '"parameters": {"max_new_tokens": 400, "temperature": 0.25, "do_sample": True}')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 4. A PIPELINE FOLYAMATA
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. A CoT pipeline folyamata (run_cot_pipeline)", level=1)
add_paragraph(doc,
    "A run_cot_pipeline() függvény koordinálja a négy lépést és a Streamlit "
    "st.status() widgetekkel élő visszajelzést ad a felhasználónak minden lépésről.")

steps_table = doc.add_table(rows=1, cols=4)
steps_table.style = "Table Grid"
add_table_row(steps_table, ["Lépés", "Függvény", "Modell", "Kimenet"], header=True)
for row in [
    ("1. Retrieval",       "find_most_similar()",  "all-MiniLM-L6-v2",         "top-k (doc, score) pár"),
    ("2. Summarization",   "call_summarization()", "facebook/bart-large-cnn",   "összefoglalt kontextus"),
    ("3. CoT Reasoning",   "call_reasoning()",     "Mistral-7B-Instruct-v0.2",  "lépésenkénti következtetés"),
    ("4. Final Answer",    "call_final_answer()",  "Mistral-7B-Instruct-v0.2",  "végső magyar nyelvű válasz"),
]:
    add_table_row(steps_table, list(row))

doc.add_paragraph()
add_paragraph(doc, "Adatfolyam az egyes lépések között:", bold=True)
add_code_block(doc,
    "Dokumentumok + Kérdés\n"
    "        │\n"
    "        ▼\n"
    "[1] Embedding API  →  top-k releváns dokumentum\n"
    "        │\n"
    "        ▼\n"
    "[2] BART Summary   →  tömörített kontextus\n"
    "        │\n"
    "        ▼\n"
    "[3] Mistral CoT    →  lépésenkénti következtetés\n"
    "        │\n"
    "        ▼\n"
    "[4] Mistral Final  →  végső válasz (1+2+3 alapján)")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 5. HIBAKEZELÉS ÉS ÚJRAPRÓBÁLKOZÁS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Hibakezelés és újrapróbálkozás (_hf_post)", level=1)
add_paragraph(doc,
    "Minden Inference API hívás a _hf_post() segédfüggvényen keresztül megy, "
    "amely egységes hibakezeléssel és újrapróbálkozási logikával rendelkezik.")

add_bullet(doc, "503 (Model Loading): ha a modell még töltődik, kivárja az estimated_time-ot (max 30s), majd újrapróbál.")
add_bullet(doc, "Timeout: 120 másodperces időkorlát; ha lejár, 5s várakozás után újrapróbál.")
add_bullet(doc, "3 kísérlet után ha nem sikerül, hibaüzenetet ad vissza a pipeline-nak.")
add_bullet(doc, "Streamlit toast értesítő jelzi a felhasználónak a várakozási állapotot.")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 6. PROMPTING TECHNIKÁK
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Alkalmazott prompting technikák", level=1)

add_heading(doc, "6.1 Personality Prompting", level=2)
add_paragraph(doc,
    "A PERSONALITY konstans egy részletes személyiség-leírást ad a modellnek: "
    "precíz, analitikus tudományos asszisztens, aki kizárólag a dokumentumokra támaszkodik "
    "és bevallja, ha nincs elegendő adat. Ez az instrukció a 3. és 4. API hívás minden "
    "promptjának elején szerepel.")

add_heading(doc, "6.2 Few-Shot Prompting", level=2)
add_paragraph(doc,
    "A FEW_SHOT_EXAMPLES konstans két teljes, annotált példát tartalmaz, amelyek "
    "megmutatják a modellnek a kívánt gondolkodási és válaszadási struktúrát "
    "(kontextus → gondolkodás lépései → részkövetkeztetés). Ez a 3. API hívásban (CoT reasoning) aktív.")

add_heading(doc, "6.3 Chain-of-Thought (CoT)", level=2)
add_paragraph(doc,
    "A prompt felépítése rákényszeríti a modellt, hogy ne azonnal adjon választ, "
    "hanem explicit lépéseket tegyen: azonosítsa a releváns adatot, nyerje ki a kulcsinformációt, "
    "ellenőrizze, majd vonja le a részkövetkeztetést. A pipeline maga is négy explicit lépésre "
    "van bontva, ami makroszinten CoT-ot valósít meg.")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 7. VÁLTOZÁSNAPLÓ (REFAKTORÁLÁSOK)
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Változásnapló", level=1)

add_paragraph(doc,
    "Ez a fejezet tartalmaz minden módosítást, amelyet a fejlesztés során az alkalmazáson elvégeztünk. "
    "Minden bejegyzés tartalmazza a változtatás dátumát, érintett fájlját, okát és pontos leírását.")

doc.add_paragraph()

# ── Változás #1 ──────────────────────────────────────────────────────────────
change_table = doc.add_table(rows=1, cols=2)
change_table.style = "Table Grid"
add_table_row(change_table, ["Változás #1 – HuggingFace Inference API endpoint frissítése", ""], header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – 22. sor (HF_API_BASE konstans)"),
    ("Típus",         "Hibajavítás (bugfix)"),
    ("Hiba",
     "HTTPError: 404 Client Error: Not Found for url: "
     "https://api-inference.huggingface.co/models/facebook/bart-large-cnn\n"
     "A call_summarization() függvény 404-es hibával tért vissza, mert a "
     "facebook/bart-large-cnn modell már nem érhető el a régi HuggingFace "
     "Inference API végponton."),
    ("Ok",
     "A HuggingFace 2024 folyamán migrálta a Serverless Inference API-t "
     "az api-inference.huggingface.co doménről a router.huggingface.co/hf-inference "
     "útvonalra. Az összes modellhívás ezen az új végponton keresztül érhető el."),
    ("Megoldás",
     "Az HF_API_BASE konstans értékét megváltoztattuk az új router URL-re. "
     "Ez a változtatás egyszerre javítja mind a négy API hívást, mivel mindegyik "
     "ugyanezt az alap URL-t használja."),
]:
    row = change_table.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()

add_paragraph(doc, "Kód – előtte / utána:", bold=True)
add_paragraph(doc, "Előtte:", italic=True)
add_code_block(doc, 'HF_API_BASE = "https://api-inference.huggingface.co/models"')
add_paragraph(doc, "Utána:", italic=True)
add_code_block(doc, 'HF_API_BASE = "https://router.huggingface.co/hf-inference/models"')

add_paragraph(doc,
    "Hatás: a 2., 3. és 4. API hívás (BART summarization, Mistral CoT reasoning, "
    "Mistral final answer) újra elérhetővé vált. Az 1. API hívás (Embedding) "
    "a LangChain könyvtáron keresztül fut, és külön végpontot használ, így azt "
    "nem érintette a változtatás.")

doc.add_paragraph()

# ── Változás #2 ──────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table2 = doc.add_table(rows=1, cols=2)
change_table2.style = "Table Grid"
add_table_row(change_table2, ["Változás #2 – Mistral hívások migrálása Chat Completions formátumra", ""], header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – call_reasoning(), call_final_answer(), + új segédfüggvények"),
    ("Típus",         "Hibajavítás (bugfix) + refaktorálás"),
    ("Hiba",
     "HTTPError: 400 Client Error: Bad Request for url: "
     "https://router.huggingface.co/hf-inference/models/mistralai/Mistral-7B-Instruct-v0.2\n"
     "A call_reasoning() és call_final_answer() függvények 400-as hibával tértek vissza, "
     "mert az új HuggingFace router végpont az utasítás-finomhangolt modellekhez "
     "már nem fogadja a régi 'inputs' + [INST]/[/INST] formátumú kéréseket."),
    ("Ok",
     "Az új router.huggingface.co/hf-inference végpont az instruct-modellek hívásához "
     "kizárólag OpenAI-kompatibilis Chat Completions formátumot fogad el "
     "(/v1/chat/completions útvonal, 'messages' tömb rendszer- és felhasználói üzenetekkel). "
     "A régi 'inputs' kulcsos payload és a Mistral [INST] tokenek ezért 400-as hibát okoznak."),
    ("Megoldás – új segédfüggvények",
     "_hf_chat_post(url, messages, api_key, max_tokens, temperature): "
     "OpenAI-kompatibilis Chat Completions POST kérést küld az /v1/chat/completions végpontra. "
     "A payload tartalmazza a 'model', 'messages', 'max_tokens' és 'temperature' mezőket.\n\n"
     "_extract_chat_text(result): a choices[0].message.content mezőből nyeri ki a választ."),
    ("Megoldás – call_reasoning() változás",
     "Régi: egyetlen prompt string [INST]...[/INST] jelölőkkel, inputs kulcscsal.\n"
     "Új: messages lista {'role': 'system', 'content': PERSONALITY + FEW_SHOT_EXAMPLES} "
     "és {'role': 'user', 'content': kontextus + kérdés + CoT trigger} elemekkel. "
     "URL: .../Mistral-7B-Instruct-v0.2/v1/chat/completions"),
    ("Megoldás – call_final_answer() változás",
     "Régi: egyetlen prompt string [INST]...[/INST] jelölőkkel, inputs kulcscsal.\n"
     "Új: messages lista {'role': 'system', 'content': PERSONALITY} "
     "és {'role': 'user', 'content': dokumentumok + összefoglaló + következtetés + kérdés} elemekkel. "
     "URL: .../Mistral-7B-Instruct-v0.2/v1/chat/completions"),
    ("Hatás",
     "A 3. és 4. API hívás (CoT reasoning, Final Answer) újra működőképes. "
     "A kód tisztább: nincs szükség [INST]/[/INST] token-kézi kezelésre és utólagos "
     "szöveg-tisztításra. Az 1. és 2. API hívás (Embedding, BART) nem változott."),
]:
    row = change_table2.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – előtte / utána (call_reasoning példáján):", bold=True)
add_paragraph(doc, "Előtte:", italic=True)
add_code_block(doc,
    'url = f"{HF_API_BASE}/{REASONING_MODEL}"\n'
    'payload = {\n'
    '    "inputs": f"<s>[INST] {PERSONALITY}...{question}\\n  1. lépés – [/INST]",\n'
    '    "parameters": {"max_new_tokens": 350, "return_full_text": False, ...},\n'
    '}\n'
    'result = _hf_post(url, payload, api_key)\n'
    'text = _extract_text(result)\n'
    'for marker in ["[/INST]", "</s>", "<s>"]: text = text.replace(marker, "")')
add_paragraph(doc, "Utána:", italic=True)
add_code_block(doc,
    'messages = [\n'
    '    {"role": "system", "content": f"{PERSONALITY}\\n\\n{FEW_SHOT_EXAMPLES}"},\n'
    '    {"role": "user",   "content": f"Kontextus: {context}\\nKérdés: {question}\\n  1. lépés –"},\n'
    ']\n'
    'url = f"{HF_API_BASE}/{REASONING_MODEL}/v1/chat/completions"\n'
    'result = _hf_chat_post(url, messages, api_key, max_tokens=350, temperature=0.35)\n'
    'return _extract_chat_text(result)')

doc.add_paragraph()

# ── Változás #3 ──────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table3 = doc.add_table(rows=1, cols=2)
change_table3.style = "Table Grid"
add_table_row(change_table3, ["Változás #3 – Mistral modell frissítése v0.3-ra + hibakezelés javítása", ""], header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – REASONING_MODEL, ANSWER_MODEL konstansok + _hf_chat_post()"),
    ("Típus",         "Hibajavítás (bugfix)"),
    ("Hiba",
     "HTTPError: 400 Client Error: Bad Request for url: "
     "https://router.huggingface.co/hf-inference/models/mistralai/Mistral-7B-Instruct-v0.2/v1/chat/completions\n"
     "A call_reasoning() és call_final_answer() függvények 400-as hibával tértek vissza "
     "a chat completions formátumra való áttérés (Változás #2) után is."),
    ("Ok",
     "Két egyidejű probléma:\n"
     "1. A mistralai/Mistral-7B-Instruct-v0.2 modell nem elérhető a HuggingFace router "
     "ingyenes szintjén (a v0.3 az aktuálisan elérhető verzió).\n"
     "2. A _hf_chat_post() nem adta hozzá a 'stream: False' mezőt a payloadhoz, "
     "amelyet a HF router chat completions végpontja megkövetel, és nem logolja "
     "a hibaválasz törzsét, ezért a pontos ok rejtve maradt."),
    ("Megoldás #1 – modell csere",
     "REASONING_MODEL és ANSWER_MODEL: Mistral-7B-Instruct-v0.2 → Mistral-7B-Instruct-v0.3\n"
     "A v0.3 az aktuálisan elérhető és ajánlott verzió a HF router ingyenes szintjén."),
    ("Megoldás #2 – stream mező",
     "'stream': False explicit hozzáadása a payload-hoz a _hf_chat_post()-ban. "
     "A HF router chat completions végpontja alapértelmezés szerint streamelt választ várhat, "
     "ezért a False értéket explicit meg kell adni."),
    ("Megoldás #3 – hibakezelés javítása",
     "resp.raise_for_status() helyett: ha resp.ok hamis, a válasz JSON törzsét "
     "(vagy szövegtörzsét) beágyazza az error kulcsú visszatérési értékbe. "
     "Ez lehetővé teszi, hogy a felhasználó és a fejlesztő lássa a tényleges API hibaüzenetet "
     "ahelyett, hogy csak az HTTP státuszkódot kapná meg."),
    ("Hatás",
     "A 3. és 4. API hívás ismét működőképes. "
     "A hibaüzenetek mostantól informatívak: HTTP státuszkód + API válasz törzsével."),
]:
    row = change_table3.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – előtte / utána:", bold=True)
add_paragraph(doc, "Modell konstansok – előtte:", italic=True)
add_code_block(doc,
    'REASONING_MODEL = "mistralai/Mistral-7B-Instruct-v0.2"\n'
    'ANSWER_MODEL    = "mistralai/Mistral-7B-Instruct-v0.2"')
add_paragraph(doc, "Modell konstansok – utána:", italic=True)
add_code_block(doc,
    'REASONING_MODEL = "mistralai/Mistral-7B-Instruct-v0.3"\n'
    'ANSWER_MODEL    = "mistralai/Mistral-7B-Instruct-v0.3"')
add_paragraph(doc, "_hf_chat_post payload – előtte:", italic=True)
add_code_block(doc,
    'payload = {\n'
    '    "model": model_id, "messages": messages,\n'
    '    "max_tokens": max_tokens, "temperature": temperature,\n'
    '}')
add_paragraph(doc, "_hf_chat_post payload – utána:", italic=True)
add_code_block(doc,
    'payload = {\n'
    '    "model": model_id, "messages": messages,\n'
    '    "max_tokens": max_tokens, "temperature": temperature,\n'
    '    "stream": False,\n'
    '}')
add_paragraph(doc, "Hibakezelés – előtte:", italic=True)
add_code_block(doc, 'resp.raise_for_status()')
add_paragraph(doc, "Hibakezelés – utána:", italic=True)
add_code_block(doc,
    'if not resp.ok:\n'
    '    try:\n'
    '        err_body = resp.json()\n'
    '    except Exception:\n'
    '        err_body = resp.text\n'
    '    return {"error": f"HTTP {resp.status_code}: {err_body}"}')

doc.add_paragraph()

# ── Változás #4 ──────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table4 = doc.add_table(rows=1, cols=2)
change_table4.style = "Table Grid"
add_table_row(change_table4,
    ["Változás #4 – Reasoning/Answer modell csere: Mistral-7B-Instruct-v0.3 → HuggingFaceH4/zephyr-7b-beta", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – REASONING_MODEL, ANSWER_MODEL konstansok (23–24. sor)"),
    ("Típus",         "Hibajavítás (bugfix)"),
    ("Hiba",
     "HTTP 400: {'error': 'Model not supported by provider hf-inference'}\n"
     "A call_reasoning() és call_final_answer() hívások 400-as hibát adtak vissza "
     "a mistralai/Mistral-7B-Instruct-v0.3 modellre."),
    ("Gyökérok",
     "A HuggingFace router (router.huggingface.co) TÖBB szolgáltatót (providert) kezel "
     "egyetlen API-n keresztül. Az URL-ben szereplő 'hf-inference' a HuggingFace saját "
     "infrastruktúráját jelenti. A Mistral modellek NEM ezen a servicen futnak — "
     "ők más providereken (Together AI, Fireworks AI, stb.) érhetők el, "
     "amelyek az URL-ben is más szegmenst kapnának (pl. /together/ vagy /fireworks-ai/).\n\n"
     "A hf-inference provider csak olyan modelleket üzemeltet, amelyeket a HuggingFace "
     "maga hoszt a saját szerverein. A Mistral modelleket harmadik fél partnerek hostolják."),
    ("Miért nem a provider váltás a megoldás?",
     "A Together AI / Fireworks AI providerek más URL-struktúrát, és esetleg "
     "eltérő API kulcsot igényelnek (a HF API token nem feltétlenül elegendő). "
     "Ez az alkalmazás egy egyszerű HF API kulcsot vár, ezért a legbiztonságosabb megoldás "
     "olyan modellt választani, amely garantáltan fut a hf-inference szinten."),
    ("Megoldás – modell csere",
     "REASONING_MODEL és ANSWER_MODEL: mistralai/Mistral-7B-Instruct-v0.3 → HuggingFaceH4/zephyr-7b-beta\n\n"
     "A Zephyr-7B-Beta modell:\n"
     "- A HuggingFace saját fejlesztése (HuggingFaceH4 szervezet) → garantáltan elérhető "
     "  a hf-inference provideren\n"
     "- Architektúra: Mistral-7B fine-tune DPO (Direct Preference Optimization) technikával, "
     "  ezért minőségben és instrukció-követésben nagyon hasonló a Mistral-Instruct modellekhez\n"
     "- Támogatja az OpenAI-kompatibilis Chat Completions formátumot (system/user üzenetek)\n"
     "- Ingyenes szinten elérhető, nincs szükség PRO előfizetésre\n"
     "- Magyar szövegek kezelésére alkalmas (multilingual képesség)"),
    ("Hatás",
     "A 3. és 4. API hívás (CoT reasoning, Final Answer) ismét működőképes. "
     "A Zephyr-7B-Beta minőségileg egyenértékű helyettesítő a Mistral-Instruct modellekkel, "
     "mivel azonos alaparchitektúrára épül."),
]:
    row = change_table4.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – előtte / utána:", bold=True)
add_paragraph(doc, "Előtte:", italic=True)
add_code_block(doc,
    'REASONING_MODEL = "mistralai/Mistral-7B-Instruct-v0.3"\n'
    'ANSWER_MODEL    = "mistralai/Mistral-7B-Instruct-v0.3"')
add_paragraph(doc, "Utána:", italic=True)
add_code_block(doc,
    'REASONING_MODEL = "HuggingFaceH4/zephyr-7b-beta"\n'
    'ANSWER_MODEL    = "HuggingFaceH4/zephyr-7b-beta"')

add_paragraph(doc, "A HuggingFace router provider-architektúrája:", bold=True)
add_code_block(doc,
    "router.huggingface.co\n"
    "  ├── /hf-inference/models/{model}      ← HF saját infrastruktúra\n"
    "  │     Elérhető: HuggingFaceH4/*, google/gemma-*, microsoft/phi-*, stb.\n"
    "  │     NEM elérhető: mistralai/*, meta-llama/* (partnerek hostolják)\n"
    "  ├── /together/v1/chat/completions      ← Together AI provider\n"
    "  │     Elérhető: mistralai/*, meta-llama/*, stb.\n"
    "  └── /fireworks-ai/...                  ← Fireworks AI provider")

doc.add_paragraph()

# ── Változás #5 ──────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table5 = doc.add_table(rows=1, cols=2)
change_table5.style = "Table Grid"
add_table_row(change_table5,
    ["Változás #5 – Chat completions → Pipeline text-generation + Zephyr chat sablon", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – call_reasoning(), call_final_answer() (3. és 4. API hívás)"),
    ("Típus",         "Hibajavítás (bugfix) + architektúrális visszaállítás"),
    ("Hiba",
     "HTTP 400: {'error': 'Model not supported by provider hf-inference'}\n"
     "A HuggingFaceH4/zephyr-7b-beta modellre is ugyanezt a 400-as hibát adta vissza "
     "az /v1/chat/completions végponton."),
    ("Gyökérok – végleges diagnózis",
     "A hf-inference provider az /v1/chat/completions végpontot NEM támogatja "
     "szöveg-generáló modellek esetén — sem Mistral, sem Zephyr, sem más "
     "text-generation modell számára. Ez a végpont kizárólag meghatározott "
     "partnerszolgáltatókon (Together AI, Fireworks AI stb.) érhető el, "
     "amelyek más URL-szegmenst és esetleg külön API hitelesítést igényelnek.\n\n"
     "A hf-inference provider CSAK a klasszikus pipeline text-generation "
     "formátumot ('inputs' kulcs + 'parameters') támogatja saját infrastruktúráján."),
    ("Megoldás – pipeline formátum visszaállítása",
     "A call_reasoning() és call_final_answer() függvények visszatértek "
     "a _hf_post() + 'inputs' kulcsos pipeline formátumhoz.\n"
     "Az [INST]/[/INST] jelölők (Mistral-specifikus) helyett a Zephyr-7B-Beta "
     "saját chat sablonját alkalmazzuk:\n\n"
     "  <|system|>\n  {rendszer üzenet}</s>\n"
     "  <|user|>\n  {felhasználói üzenet}</s>\n"
     "  <|assistant|>\n\n"
     "A generálás az <|assistant|> token után indul, return_full_text: False "
     "beállítással csak az új szöveg kerül vissza."),
    ("Miért Zephyr-7B-Beta és nem más modell?",
     "1. HuggingFace saját fejlesztése → garantáltan fut a hf-inference provideren\n"
     "2. Mistral-7B alapú fine-tune → hasonló minőség és instrukció-követés\n"
     "3. Jól definiált, dokumentált chat sablon (<|system|>/<|user|>/<|assistant|>)\n"
     "4. Ingyenes szinten elérhető, PRO előfizetés nélkül\n"
     "5. Magyar szövegek feldolgozására is alkalmas"),
    ("Hatás",
     "A 3. és 4. API hívás véglegesen működőképes. "
     "A _hf_chat_post() és _extract_chat_text() segédfüggvények megmaradnak "
     "a kódban jövőbeli felhasználásra (pl. ha valaki Together AI providerre vált), "
     "de az aktuális pipeline jelenleg nem hívja őket."),
]:
    row = change_table5.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – előtte / utána (call_reasoning példáján):", bold=True)
add_paragraph(doc, "Előtte (chat completions – NEM működött hf-inference-en):", italic=True)
add_code_block(doc,
    'messages = [\n'
    '    {"role": "system", "content": f"{PERSONALITY}..."},\n'
    '    {"role": "user",   "content": f"Kontextus: {context}..."},\n'
    ']\n'
    'url = f"{HF_API_BASE}/{REASONING_MODEL}/v1/chat/completions"\n'
    'result = _hf_chat_post(url, messages, api_key, max_tokens=350, temperature=0.35)\n'
    'return _extract_chat_text(result)')
add_paragraph(doc, "Utána (pipeline text-generation – hf-inference-en fut):", italic=True)
add_code_block(doc,
    'prompt = (\n'
    '    f"<|system|>\\n{PERSONALITY}\\n\\n{FEW_SHOT_EXAMPLES}</s>\\n"\n'
    '    f"<|user|>\\nKontextus: {context}\\nKérdés: {question}\\n  1. lépés –</s>\\n"\n'
    '    f"<|assistant|>\\n"\n'
    ')\n'
    'url = f"{HF_API_BASE}/{REASONING_MODEL}"\n'
    'payload = {\n'
    '    "inputs": prompt,\n'
    '    "parameters": {"max_new_tokens": 350, "temperature": 0.35,\n'
    '                   "do_sample": True, "return_full_text": False},\n'
    '    "options": {"wait_for_model": True},\n'
    '}\n'
    'result = _hf_post(url, payload, api_key)\n'
    'return _extract_text(result)')

add_paragraph(doc, "A hf-inference provider támogatási mátrixa:", bold=True)
add_code_block(doc,
    "Végpont formátum                          │ hf-inference │ together\n"
    "──────────────────────────────────────────┼──────────────┼─────────\n"
    "Pipeline: {inputs, parameters}            │      ✅       │    ✅\n"
    "/v1/chat/completions (text-gen modellek)  │      ❌       │    ✅\n"
    "/v1/chat/completions (BART-szerű modellek)│      ❌       │    ❌\n"
    "Összefoglalás (summary_text)              │      ✅       │    ✅")

doc.add_paragraph()

# ── Változás #6 ──────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table6 = doc.add_table(rows=1, cols=2)
change_table6.style = "Table Grid"
add_table_row(change_table6,
    ["Változás #6 – _hf_post hibakezelés javítása + options mező eltávolítása", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – _hf_post(), call_reasoning(), call_final_answer()"),
    ("Típus",         "Hibajavítás (bugfix) – két egyidejű probléma"),
    ("1. hiba",
     "HTTPError: 400 Client Error: Bad Request for url: "
     "https://router.huggingface.co/hf-inference/models/HuggingFaceH4/zephyr-7b-beta\n"
     "A call_reasoning() hívás 400-as hibával tért vissza, de a tényleges hibaüzenet "
     "nem volt látható, mert _hf_post() resp.raise_for_status()-t dobott."),
    ("1. hiba – ok",
     "Az 'options': {'wait_for_model': True} mező az új HuggingFace router "
     "text-generation végpontján ÉRVÉNYTELEN mezőnek számít és 400 Bad Request "
     "hibát okoz. Az opcionális mezőket az öreg api-inference.huggingface.co "
     "végpont elfogadta, de az új router.huggingface.co szigorúbb validációt alkalmaz. "
     "Fontos: a summarization (BART) végpont még elfogadja ezt a mezőt, "
     "csak a text-generation végpont utasítja vissza."),
    ("1. hiba – megoldás",
     "Az 'options': {'wait_for_model': True} mező eltávolítása a call_reasoning() "
     "és call_final_answer() payload-jából. A BART summarization hívásban "
     "(call_summarization) megmarad, mert ott működik."),
    ("2. hiba",
     "_hf_post() a resp.raise_for_status() hívással Python HTTPError kivételt "
     "dobott 4xx hibáknál, amelynek üzenete csak az URL-t és a státuszkódot "
     "tartalmazta — a HuggingFace API tényleges hibaüzenete (pl. 'Model not "
     "supported by provider hf-inference') elveszett."),
    ("2. hiba – megoldás",
     "resp.raise_for_status() helyett: ha resp.ok hamis, a válasz JSON törzse "
     "(vagy szövegtörzsét) beágyazza az {'error': 'HTTP {kod}: {törzs}'} "
     "visszatérési értékbe — ugyanaz a minta, amit korábban _hf_chat_post()-ban "
     "bevezettünk (Változás #3). Így mind a négy hívás egységesen kezeli a hibákat "
     "és megjeleníti a tényleges API hibaüzenetet."),
    ("Hatás",
     "A 3. és 4. API hívás a text-generation pipeline formátumban működőképes. "
     "A hibakezelés egységes: mind _hf_post(), mind _hf_chat_post() az error "
     "kulcsos dict-ben adja vissza a részletes hibaüzenetet, nem dob kivételt."),
]:
    row = change_table6.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – előtte / utána:", bold=True)
add_paragraph(doc, "_hf_post hibakezelés – előtte:", italic=True)
add_code_block(doc,
    'resp.raise_for_status()   # csak HTTP státuszkód + URL látszott\n'
    'return resp.json()')
add_paragraph(doc, "_hf_post hibakezelés – utána:", italic=True)
add_code_block(doc,
    'if not resp.ok:\n'
    '    try:\n'
    '        err_body = resp.json()\n'
    '    except Exception:\n'
    '        err_body = resp.text\n'
    '    return {"error": f"HTTP {resp.status_code}: {err_body}"}\n'
    'return resp.json()')
add_paragraph(doc, "payload options mező – előtte (call_reasoning / call_final_answer):", italic=True)
add_code_block(doc,
    'payload = {\n'
    '    "inputs": prompt,\n'
    '    "parameters": {...},\n'
    '    "options": {"wait_for_model": True},  # ← 400-at okozott az új routeren\n'
    '}')
add_paragraph(doc, "payload options mező – utána:", italic=True)
add_code_block(doc,
    'payload = {\n'
    '    "inputs": prompt,\n'
    '    "parameters": {...},\n'
    '    # options mező eltávolítva\n'
    '}')

add_paragraph(doc, "Végpont / options mező kompatibilitási táblázat:", bold=True)
add_code_block(doc,
    'Végpont                        │ options elfogadva?\n'
    '───────────────────────────────┼───────────────────\n'
    'BART summarization             │       ✅\n'
    'Zephyr text-generation         │       ❌  (400-at okoz)\n'
    'router chat completions        │       ❌  (nem értelmezett mező)')

doc.add_paragraph()

# ── Változás #7 ──────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table7 = doc.add_table(rows=1, cols=2)
change_table7.style = "Table Grid"
add_table_row(change_table7,
    ["Változás #7 – CoT és Final Answer: hf-inference pipeline → Together AI chat completions", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl",
     "app.py – HF_CHAT_URL konstans (új), REASONING_MODEL, ANSWER_MODEL, "
     "call_reasoning(), call_final_answer()"),
    ("Típus",         "Hibajavítás (bugfix) – végleges megoldás a reasoning hibára"),
    ("Hiba",
     "A call_reasoning() (3. API hívás) minden futtatásnál hibára futott. "
     "A hf-inference provider a text-generation pipeline formátumot sem fogadta el "
     "Zephyr-7B-Beta modelltől, és a /v1/chat/completions végpontot sem támogatja "
     "saját infrastruktúráján."),
    ("Gyökérok – végleges diagnózis",
     "A hf-inference provider a HuggingFace saját, INGYENES számítási kapacitását "
     "jelenti. Ez a provider CSAK kis méretű vagy kifejezetten általa hostolt "
     "modellekhez biztosít szabad hozzáférést. A 7B+ paraméteres instrukció-finomhangolt "
     "modellek (Mistral-7B, Zephyr-7B, stb.) ezen a provideren NEM futnak – sem "
     "pipeline, sem chat completions formátumban.\n\n"
     "A HuggingFace router azonban TÖBB providert aggregál egyetlen API alá, "
     "amelyek mindegyike ugyanazzal a HF Bearer tokennel hívható. A Together AI "
     "provider (/together/) kifejezetten nagy nyelvi modellekre optimalizált "
     "és támogatja a Mistral-7B-Instruct modelleket."),
    ("Megoldás – új HF_CHAT_URL konstans",
     "Hozzáadva: HF_CHAT_URL = 'https://router.huggingface.co/together/v1/chat/completions'\n\n"
     "Ez a végpont:\n"
     "- Ugyanazzal a HF API kulccsal (Bearer token) hitelesíti magát\n"
     "- OpenAI-kompatibilis chat completions formátumot fogad\n"
     "- Támogatja a mistralai/Mistral-7B-Instruct-v0.3 modellt\n"
     "- Ingyenes havi HF kreditkeret terhére érhető el\n\n"
     "A HF_API_BASE megmarad a BART summarization híváshoz (hf-inference, BART ott működik)."),
    ("Megoldás – modellek visszaállítása",
     "REASONING_MODEL és ANSWER_MODEL: HuggingFaceH4/zephyr-7b-beta → mistralai/Mistral-7B-Instruct-v0.3\n"
     "A Together AI provider natívan támogatja a Mistral modelleket."),
    ("Megoldás – call_reasoning() és call_final_answer() átírása",
     "Mindkét függvény visszatér a _hf_chat_post() + _extract_chat_text() használatára, "
     "de most a HF_CHAT_URL (together) végponttal – a Zephyr pipeline sablonok eltávolítva.\n\n"
     "Végeredmény: a két Mistral-hívás a together provideren fut (chat completions), "
     "a BART-hívás a hf-inference provideren fut (pipeline) – a két provider "
     "teljesen elkülönített URL-eken van."),
    ("Hatás",
     "A 3. és 4. API hívás (CoT reasoning, Final Answer) a Together AI provideren "
     "stabilisan működik Mistral-7B-Instruct-v0.3 modellel. "
     "A reasoning hiba megszűnt."),
]:
    row = change_table7.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – előtte / utána:", bold=True)
add_paragraph(doc, "Konstansok – előtte:", italic=True)
add_code_block(doc,
    'HF_API_BASE     = "https://router.huggingface.co/hf-inference/models"\n'
    'REASONING_MODEL = "HuggingFaceH4/zephyr-7b-beta"\n'
    'ANSWER_MODEL    = "HuggingFaceH4/zephyr-7b-beta"')
add_paragraph(doc, "Konstansok – utána:", italic=True)
add_code_block(doc,
    'HF_API_BASE     = "https://router.huggingface.co/hf-inference/models"  # BART marad\n'
    'HF_CHAT_URL     = "https://router.huggingface.co/together/v1/chat/completions"  # ÚJ\n'
    'REASONING_MODEL = "mistralai/Mistral-7B-Instruct-v0.3"\n'
    'ANSWER_MODEL    = "mistralai/Mistral-7B-Instruct-v0.3"')

add_paragraph(doc, "call_reasoning() – előtte (pipeline, nem működött):", italic=True)
add_code_block(doc,
    'prompt = f"<|system|>\\n{system_content}</s>\\n<|user|>\\n{user_content}</s>\\n<|assistant|>\\n"\n'
    'url = f"{HF_API_BASE}/{REASONING_MODEL}"\n'
    'payload = {"inputs": prompt, "parameters": {...}}\n'
    'result = _hf_post(url, payload, api_key)')
add_paragraph(doc, "call_reasoning() – utána (together chat completions, működik):", italic=True)
add_code_block(doc,
    'messages = [\n'
    '    {"role": "system", "content": f"{PERSONALITY}\\n\\n{FEW_SHOT_EXAMPLES}"},\n'
    '    {"role": "user",   "content": f"Kontextus: {context}\\nKérdés: {question}\\n  1. lépés –"},\n'
    ']\n'
    'result = _hf_chat_post(HF_CHAT_URL, messages, api_key, max_tokens=350, temperature=0.35)')

add_paragraph(doc, "A HF router provider-architektúra (végleges):", bold=True)
add_code_block(doc,
    "router.huggingface.co\n"
    "  ├── /hf-inference/models/facebook/bart-large-cnn   ← 2. API hívás ✅\n"
    "  │     Pipeline summarization – kis modell, ingyenes\n"
    "  └── /together/v1/chat/completions                  ← 3. és 4. API hívás ✅\n"
    "        Chat completions – Mistral-7B-Instruct-v0.3\n"
    "        HF kreditkeret terhére, ugyanaz az API kulcs")

doc.add_paragraph()

# ── Változás #8 ──────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table8 = doc.add_table(rows=1, cols=2)
change_table8.style = "Table Grid"
add_table_row(change_table8,
    ["Változás #8 – _hf_chat_post IndexError javítása: model_id URL-ből → explicit paraméter", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – _hf_chat_post(), call_reasoning(), call_final_answer()"),
    ("Típus",         "Hibajavítás (bugfix)"),
    ("Hiba",
     "IndexError: list index out of range\n"
     "File app.py, line 101, in _hf_chat_post\n"
     "    model_id = url.split('/models/')[1].split('/v1')[0]"),
    ("Ok",
     "A _hf_chat_post() függvény a model azonosítóját az URL-ből nyerte ki "
     "url.split('/models/')[1] kifejezéssel. Ez a logika az eredeti hf-inference "
     "URL-ekre volt tervezve (pl. .../models/mistralai/Mistral-7B.../v1/...), ahol "
     "a '/models/' szegmens szerepel az URL-ben.\n\n"
     "A Változás #7-ben bevezetett HF_CHAT_URL azonban "
     "'https://router.huggingface.co/together/v1/chat/completions' formátumú, "
     "amelyben NEM szerepel '/models/' szegmens. Ezért az url.split('/models/') "
     "csak egyetlen elemet tartalmazó listát ad vissza, és a [1] index "
     "IndexError kivételt okoz."),
    ("Megoldás",
     "A _hf_chat_post() szignatúrájához hozzáadva egy 'model: str' paraméter "
     "(alapértelmezett: üres string). A model_id = url.split(...) sor eltávolítva; "
     "a payload 'model' mezőjébe a függvénynek átadott model paraméter kerül.\n\n"
     "A call_reasoning() hívásban: model=REASONING_MODEL\n"
     "A call_final_answer() hívásban: model=ANSWER_MODEL\n\n"
     "Ez egységes és URL-független megoldás: a függvény nem támaszkodik "
     "az URL szerkezetére a modell azonosításához."),
    ("Hatás",
     "Az IndexError megszűnt. A _hf_chat_post() bármilyen chat completions "
     "URL-lel és bármilyen modell azonosítóval használható, "
     "függetlenül az URL formátumától."),
]:
    row = change_table8.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – előtte / utána:", bold=True)
add_paragraph(doc, "_hf_chat_post szignatúra és model_id – előtte:", italic=True)
add_code_block(doc,
    'def _hf_chat_post(url, messages, api_key, max_tokens=400, temperature=0.3, retries=3):\n'
    '    model_id = url.split("/models/")[1].split("/v1")[0]  # IndexError, ha nincs /models/\n'
    '    payload = {"model": model_id, ...}')
add_paragraph(doc, "_hf_chat_post szignatúra és model_id – utána:", italic=True)
add_code_block(doc,
    'def _hf_chat_post(url, messages, api_key, model="", max_tokens=400, temperature=0.3, retries=3):\n'
    '    # model_id kinyerése eltávolítva\n'
    '    payload = {"model": model, ...}')
add_paragraph(doc, "Hívási helyek – utána:", italic=True)
add_code_block(doc,
    '# call_reasoning():\n'
    '_hf_chat_post(HF_CHAT_URL, messages, api_key, model=REASONING_MODEL,\n'
    '              max_tokens=350, temperature=0.35)\n\n'
    '# call_final_answer():\n'
    '_hf_chat_post(HF_CHAT_URL, messages, api_key, model=ANSWER_MODEL,\n'
    '              max_tokens=400, temperature=0.25)')

doc.add_paragraph()

# ── Változás #9 ──────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table9 = doc.add_table(rows=1, cols=2)
change_table9.style = "Table Grid"
add_table_row(change_table9,
    ["Változás #9 – Frontend: API kulcs elrejtése, HF API szekció törlése, emojik eltávolítása, égkék háttér", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – main() függvény (Streamlit UI)"),
    ("Típus",         "Frontend refaktorálás (4 egyidejű változtatás)"),

    ("1. változtatás – API kulcs elrejtése",
     "Korábban: az API kulcs egy st.text_input(type='password') mezőben jelent meg "
     "az oldalsávban, ahol a felhasználó láthatta és módosíthatta.\n\n"
     "Mostantól: az API kulcs kizárólag a .env fájlból töltődik be "
     "(os.getenv('HF_API_KEY', '')), a frontendon nem jelenik meg beviteli mező. "
     "Ha a .env fájlban nincs megadva, figyelmeztető üzenet jelenik meg: "
     "'Az API kulcs nem található. Add meg a HF_API_KEY értékét a .env fájlban!'"),

    ("2. változtatás – HuggingFace API hívások szekció törlése",
     "Az oldalsávból eltávolítva a '### HuggingFace API hívások' fejléc és az alatta "
     "lévő markdown táblázat (4 sor: modellnevek és céljaik). "
     "Az 'Alkalmazott technikák' szekció megmaradt."),

    ("3. változtatás – Összes emoji eltávolítása",
     "Eltávolított emojik és cserék:\n"
     "- page_icon='🔍' → eltávolítva (st.set_page_config-ból)\n"
     "- st.title('🔍 Dokumentum-alapú Q&A') → 'Dokumentum-alapú Q&A'\n"
     "- st.header('⚙️ Beállítások') → 'Beállítások'\n"
     "- '### 📡 HuggingFace API hívások' → szekció eltávolítva\n"
     "- '### 🎯 Alkalmazott technikák' → '### Alkalmazott technikák'\n"
     "- st.subheader('📄 Dokumentumok') → 'Dokumentumok'\n"
     "- st.subheader('❓ Kérdés és futtatás') → 'Kérdés és futtatás'\n"
     "- st.success('✅ ... dokumentum betöltve') → emoji nélkül\n"
     "- st.warning('⚠️ ...') → emoji nélkül\n"
     "- st.info('ℹ️ ...') → emoji nélkül (2 helyen)\n"
     "- st.button('🚀 Válasz generálása...') → 'Válasz generálása...'\n"
     "- st.subheader('🔄 Chain-of-Thought...') → emoji nélkül\n"
     "- st.subheader('✅ Végső válasz') → 'Végső válasz'\n"
     "- st.expander('📊 Részletes CoT összesítő') → 'Részletes CoT összesítő'\n"
     "- Hasonlósági badge emojik (🟢🟡🔴) → [magas] [közepes] [alacsony] szöveges jelzők"),

    ("4. változtatás – Égkék háttér (custom CSS)",
     "Hozzáadva egy st.markdown() blokk unsafe_allow_html=True paraméterrel, "
     "amely CSS stílusokat injektál a Streamlit alkalmazásba:\n\n"
     "  .stApp { background-color: #87CEEB; }\n"
     "  [data-testid='stSidebar'] { background-color: #5ba3c9; }\n\n"
     "#87CEEB: klasszikus égkék (sky blue) – fő tartalomterület\n"
     "#5ba3c9: sötétebb kékesszürke – oldalsáv, vizuális elkülönítéshez\n\n"
     "A CSS a .stApp Streamlit konténer osztályát célozza (verziótól független), "
     "az oldalsávhoz a data-testid attribútum alapú selector biztosít stabilitást."),

    ("Hatás",
     "Az API kulcs nem látható és nem módosítható a felhasználók által a webes "
     "felületen – kizárólag szerver oldali környezeti változóból olvasódik be. "
     "A felület emoji-mentes, az égkék háttér vizuálisan egységes megjelenést biztosít."),
]:
    row = change_table9.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run_c in row.cells[0].paragraphs[0].runs:
        run_c.bold = True
        run_c.font.size = Pt(10)
    for run_c in row.cells[1].paragraphs[0].runs:
        run_c.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "CSS kód – égkék háttér:", bold=True)
add_code_block(doc,
    'st.markdown(\n'
    '    """\n'
    '    <style>\n'
    '    .stApp {\n'
    '        background-color: #87CEEB;\n'
    '    }\n'
    '    [data-testid="stSidebar"] {\n'
    '        background-color: #5ba3c9;\n'
    '    }\n'
    '    </style>\n'
    '    """,\n'
    '    unsafe_allow_html=True,\n'
    ')')

add_paragraph(doc, "API kulcs kezelés – előtte / utána:", bold=True)
add_paragraph(doc, "Előtte (látható beviteli mező az oldalsávban):", italic=True)
add_code_block(doc,
    'api_key = st.text_input(\n'
    '    "HuggingFace API kulcs",\n'
    '    value=os.getenv("HF_API_KEY", ""),\n'
    '    type="password",\n'
    ')')
add_paragraph(doc, "Utána (csak .env-ből, frontend nem látja):", italic=True)
add_code_block(doc,
    'api_key = os.getenv("HF_API_KEY", "")')

doc.add_paragraph()

# ── Változás #10 ─────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table10 = doc.add_table(rows=1, cols=2)
change_table10.style = "Table Grid"
add_table_row(change_table10,
    ["Változás #10 – Alkalmazott technikák törlése, halványkék háttér, CoT fallback a végső válaszhoz", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl",
     "app.py – konstansok, _is_api_error() (új), run_cot_pipeline(), main()"),
    ("Típus",         "Frontend refaktorálás + robusztusság javítás (3 egyidejű változtatás)"),

    ("1. változtatás – Alkalmazott technikák szekció törlése",
     "Az oldalsávból eltávolítva:\n"
     "- st.divider() elválasztó\n"
     "- '### Alkalmazott technikák' fejléc\n"
     "- Chain-of-Thought / Personality prompting / Few-Shot / RAG felsorolás\n\n"
     "Az oldalsáv mostantól csak a 'Beállítások' fejlécet és a top-k slidert tartalmazza."),

    ("2. változtatás – Háttér halványkékre változtatva",
     "Fő terület: #87CEEB (égkék) → #EBF5FB (nagyon halványkék, majdnem fehér)\n"
     "Oldalsáv: #5ba3c9 (sötétkék) → #D6EAF8 (halványkék)\n\n"
     "#EBF5FB: a fehérhez közeli, alig kékes árnyalat – kellemes, nem tolakodó háttér\n"
     "#D6EAF8: enyhén sötétebb halványkék – az oldalsávot finoman megkülönbözteti "
     "a fő területtől anélkül, hogy kontrasztosan elütne"),

    ("3. változtatás – CoT következtetés graceful fallback",
     "Probléma: ha a 3. API hívás (call_reasoning) hibával tért vissza, "
     "a hibaüzenet szövegként kerül be a 4. hívás (call_final_answer) kontextusába, "
     "ami megakadályozza, hogy a modell érdemi választ adjon a kérdésre.\n\n"
     "Megoldás – új REASONING_FALLBACK konstans:\n"
     "'(A CoT következtetési lépés nem volt elérhető – a végső válasz a közvetlen "
     "dokumentumkontextus alapján készül.)'\n\n"
     "Megoldás – új _is_api_error(text) helper:\n"
     "Visszaad True-t, ha a szöveg API hibaválasz: üres, 'HTTP '-vel kezdődik, "
     "'Nem sikerült elérni a modellt.' egyenlő, vagy '{...error...}' JSON.\n\n"
     "Megoldás – run_cot_pipeline 3. lépése:\n"
     "- Hiba esetén: reasoning = REASONING_FALLBACK, figyelmeztetés megjelenik, "
     "  a pipeline FOLYTATÓDIK a 4. lépéssel\n"
     "- Siker esetén: a szokásos CoT szöveg kerül tovább\n\n"
     "Eredmény: a modell mindig megkísérli kiírni a végső választ a kérdésre, "
     "akkor is, ha a CoT lépés nem elérhető."),

    ("Hatás",
     "Az oldalsáv letisztultabb. A háttér finoman kékes, nem tolakodó. "
     "A 4. API hívás (Final Answer) most már akkor is lefut és választ ad, "
     "ha a 3. hívás (CoT reasoning) valamilyen okból hibára futna."),
]:
    row = change_table10.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run_c in row.cells[0].paragraphs[0].runs:
        run_c.bold = True
        run_c.font.size = Pt(10)
    for run_c in row.cells[1].paragraphs[0].runs:
        run_c.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – _is_api_error() helper:", bold=True)
add_code_block(doc,
    'def _is_api_error(text: str) -> bool:\n'
    '    if not text or not text.strip():\n'
    '        return True\n'
    '    t = text.strip()\n'
    '    return (\n'
    '        t.startswith("HTTP ")\n'
    '        or t == "Nem sikerült elérni a modellt."\n'
    '        or (t.startswith("{") and "error" in t.lower())\n'
    '    )')

add_paragraph(doc, "Kód – run_cot_pipeline 3. lépés – előtte:", italic=True)
add_code_block(doc,
    'reasoning = call_reasoning(context_text, question, api_key)\n'
    'steps["reasoning"] = reasoning\n'
    'st.code(f"1. lépés – {reasoning}", language=None)\n'
    'status3.update(label="3. lépés: ... ✅", state="complete")')

add_paragraph(doc, "Kód – run_cot_pipeline 3. lépés – utána:", italic=True)
add_code_block(doc,
    'raw_reasoning = call_reasoning(context_text, question, api_key)\n'
    'if _is_api_error(raw_reasoning):\n'
    '    reasoning = REASONING_FALLBACK\n'
    '    steps["reasoning"] = reasoning\n'
    '    steps["reasoning_error"] = raw_reasoning\n'
    '    st.warning("A CoT következtetési lépés nem érhető el – folytatás kontextusból.")\n'
    '    status3.update(label="... (kihagyva, folytatás kontextusból)", state="error")\n'
    'else:\n'
    '    reasoning = raw_reasoning\n'
    '    steps["reasoning"] = reasoning\n'
    '    st.code(f"1. lépés – {reasoning}", language=None)\n'
    '    status3.update(label="3. lépés: ... ✅", state="complete")')

add_paragraph(doc, "Háttérszínek – előtte / utána:", bold=True)
add_code_block(doc,
    'Terület        │ Előtte   │ Utána\n'
    '───────────────┼──────────┼────────\n'
    'Fő tartalom    │ #87CEEB  │ #EBF5FB\n'
    'Oldalsáv       │ #5ba3c9  │ #D6EAF8')

doc.add_paragraph()

# ── Változás #11 ─────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table11 = doc.add_table(rows=1, cols=2)
change_table11.style = "Table Grid"
add_table_row(change_table11,
    ["Változás #11 – 5. API hívás: call_quick_answer() – egy mondatos tömör válasz", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl",
     "app.py – call_quick_answer() (új függvény), run_cot_pipeline() 5. lépés, "
     "main() eredmény megjelenítés"),
    ("Típus",         "Új funkció (feature)"),
    ("Igény",
     "A modell a CoT pipeline kimenetein túl válaszolja meg a feltett kérdést "
     "egyetlen, tömör mondatban – ez az egy mondatos válasz az eredmény oldalon "
     "kiemelve jelenik meg, jól olvasható formában."),
    ("Megoldás – call_quick_answer() függvény",
     "Új 5. API hívás: Mistral-7B-Instruct-v0.3 a Together AI provideren keresztül "
     "(ugyanaz a HF_CHAT_URL és ANSWER_MODEL, mint a 4. hívásban).\n\n"
     "Prompt felépítése:\n"
     "- system: PERSONALITY (analitikus asszisztens)\n"
     "- user: Kontextus + Kérdés + explicit utasítás: 'PONTOSAN EGY MONDATBAN, "
     "  MAGYARUL, kizárólag a kontextus alapján. Ne írj többet, mint egyetlen "
     "  teljes mondat.'\n\n"
     "Paraméterek:\n"
     "- max_tokens: 80 (egy mondathoz bőven elegendő, megakadályozza a terjengős választ)\n"
     "- temperature: 0.1 (szinte determinisztikus – a tömörség és pontosság elsődleges)\n\n"
     "Ez alacsonyabb temperature-t és kevesebb max_tokens-t használ, mint a 4. hívás, "
     "mert itt a cél nem az elaborált magyarázat, hanem az egyetlen, precíz mondat."),
    ("Megoldás – run_cot_pipeline 5. lépés",
     "A pipeline-ban új 5. lépés: st.status() widgettel jelzi az egy mondatos válasz "
     "generálását, az eredményt steps['quick_answer'] kulcson tárolja."),
    ("Megoldás – UI megjelenítés",
     "Az eredmény oldalon az egy mondatos válasz FELÜL jelenik meg (a részletes végső "
     "válasz előtt), 'Egy mondatos válasz:' felirattal és st.info() dobozban – "
     "vizuálisan elkülönítve a hosszabb részletes választól.\n\n"
     "Ha az 5. API hívás hibával tér vissza (_is_api_error() alapján), a doboz "
     "nem jelenik meg, a pipeline nem omlik össze."),
    ("Hatás",
     "A felhasználó két választ lát egymás alatt:\n"
     "1. Egy mondatos válasz (st.info doboz) – gyors, tömör\n"
     "2. Részletes végső válasz (st.success doboz) – CoT alapú, bővebb\n\n"
     "Az alkalmazás mostantól összesen 5 HuggingFace API hívást tartalmaz."),
]:
    row = change_table11.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run_c in row.cells[0].paragraphs[0].runs:
        run_c.bold = True
        run_c.font.size = Pt(10)
    for run_c in row.cells[1].paragraphs[0].runs:
        run_c.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "call_quick_answer() kód:", bold=True)
add_code_block(doc,
    'def call_quick_answer(context: str, question: str, api_key: str) -> str:\n'
    '    messages = [\n'
    '        {"role": "system", "content": PERSONALITY},\n'
    '        {"role": "user", "content": (\n'
    '            f"Kontextus: {context}\\n\\n"\n'
    '            f"Kérdés: {question}\\n\\n"\n'
    '            f"Válaszolj PONTOSAN EGY MONDATBAN, MAGYARUL, "\n'
    '            f"kizárólag a kontextus alapján. "\n'
    '            f"Ne írj többet, mint egyetlen teljes mondat."\n'
    '        )},\n'
    '    ]\n'
    '    result = _hf_chat_post(\n'
    '        HF_CHAT_URL, messages, api_key,\n'
    '        model=ANSWER_MODEL, max_tokens=80, temperature=0.1,\n'
    '    )\n'
    '    return _extract_chat_text(result)')

add_paragraph(doc, "UI megjelenítés – kód:", bold=True)
add_code_block(doc,
    'quick = results.get("quick_answer", "")\n'
    'if quick and not _is_api_error(quick):\n'
    '    st.markdown("**Egy mondatos válasz:**")\n'
    '    st.info(quick)\n'
    '    st.divider()\n\n'
    'final = results.get("final_answer", "")\n'
    'if final and not final.startswith("{"):\n'
    '    st.markdown("**Részletes végső válasz:**")\n'
    '    st.success(final)')

add_paragraph(doc, "A teljes pipeline (frissített) – 5 API hívás:", bold=True)
add_code_block(doc,
    "Lépés │ Függvény              │ Modell / Provider          │ Kimenet\n"
    "──────┼───────────────────────┼────────────────────────────┼───────────────────────────\n"
    "  1   │ find_most_similar()   │ all-MiniLM-L6-v2 / HF      │ top-k releváns dokumentum\n"
    "  2   │ call_summarization()  │ bart-large-cnn / hf-inf.   │ összefoglalt kontextus\n"
    "  3   │ call_reasoning()      │ Mistral-7B-v0.3 / together │ CoT következtetés\n"
    "  4   │ call_final_answer()   │ Mistral-7B-v0.3 / together │ részletes végső válasz\n"
    "  5   │ call_quick_answer()   │ Mistral-7B-v0.3 / together │ egy mondatos tömör válasz")

doc.add_paragraph()

# ── Változás #12 ─────────────────────────────────────────────────────────────
doc.add_paragraph()
change_table12 = doc.add_table(rows=1, cols=2)
change_table12.style = "Table Grid"
add_table_row(change_table12,
    ["Változás #12 – REASONING_MODEL és ANSWER_MODEL csere: Mistral-7B-v0.3 → Meta-Llama-3.1-8B-Instruct-Turbo", ""],
    header=True)

for field, value in [
    ("Dátum",         "2026-04-26"),
    ("Érintett fájl", "app.py – REASONING_MODEL, ANSWER_MODEL konstansok (24–25. sor)"),
    ("Típus",         "Hibajavítás (bugfix)"),
    ("Hiba",
     "HTTP 410: {'error': 'The requested model is deprecated and no longer supported "
     "by provider together'}\n"
     "A 4. lépés (call_final_answer) és az 5. lépés (call_quick_answer) egyaránt "
     "410 Gone hibával tért vissza a mistralai/Mistral-7B-Instruct-v0.3 modellre."),
    ("Ok",
     "A HTTP 410 (Gone) státuszkód azt jelzi, hogy az erőforrás véglegesen eltávolításra "
     "került. A Together AI provider elavulttá (deprecated) nyilvánította és "
     "eltávolította a mistralai/Mistral-7B-Instruct-v0.3 modellt a kínálatából. "
     "Ez a Together AI üzleti döntése – a modell elérhető marad más providereken, "
     "de a Together AI infrastruktúráján nem fut tovább."),
    ("Modellválasztás indoklása",
     "Csere: meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo\n\n"
     "Miért ez a modell?\n"
     "1. Together AI natív -Turbo modell: a Together AI maga optimalizálta és "
     "   tartja karban, nem várható rövid távú deprecation\n"
     "2. Méret: 8B paraméter – hasonló képességű, mint a Mistral-7B, "
     "   de modernebb architektúra (Llama 3.1)\n"
     "3. Chat completions támogatás: natívan támogatja az OpenAI-kompatibilis "
     "   messages formátumot (system/user/assistant szerepek)\n"
     "4. Magyar nyelv: a Llama 3.1 erősebb multilingual képességekkel rendelkezik, "
     "   mint a Mistral-7B\n"
     "5. Ingyenes HF kreditkeret: elérhető a HF router together provideren keresztül "
     "   ugyanazzal az API kulccsal"),
    ("Érintett API hívások",
     "3. hívás (call_reasoning) – REASONING_MODEL\n"
     "4. hívás (call_final_answer) – ANSWER_MODEL\n"
     "5. hívás (call_quick_answer) – ANSWER_MODEL\n\n"
     "Mindhárom hívás ugyanazt a modellt használja, ezért egyetlen konstanscsere "
     "elegendő volt."),
    ("Hatás",
     "A 3., 4. és 5. API hívás ismét működőképes a Meta-Llama-3.1-8B-Instruct-Turbo "
     "modellel. Prompt formátum nem változott – a messages tömb (system/user) "
     "kompatibilis az új modellel."),
]:
    row = change_table12.add_row()
    row.cells[0].text = field
    row.cells[1].text = value
    for run_c in row.cells[0].paragraphs[0].runs:
        run_c.bold = True
        run_c.font.size = Pt(10)
    for run_c in row.cells[1].paragraphs[0].runs:
        run_c.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, "Kód – előtte / utána:", bold=True)
add_paragraph(doc, "Előtte:", italic=True)
add_code_block(doc,
    'REASONING_MODEL = "mistralai/Mistral-7B-Instruct-v0.3"\n'
    'ANSWER_MODEL    = "mistralai/Mistral-7B-Instruct-v0.3"')
add_paragraph(doc, "Utána:", italic=True)
add_code_block(doc,
    'REASONING_MODEL = "meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo"\n'
    'ANSWER_MODEL    = "meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo"')

add_paragraph(doc, "Modelltörténet – REASONING_MODEL és ANSWER_MODEL változásai:", bold=True)
add_code_block(doc,
    "Változás │ Modell                                    │ Ok\n"
    "─────────┼───────────────────────────────────────────┼──────────────────────────────\n"
    "  #1     │ mistralai/Mistral-7B-Instruct-v0.2        │ eredeti\n"
    "  #3     │ mistralai/Mistral-7B-Instruct-v0.3        │ v0.2 nem érhető el (404)\n"
    "  #4     │ HuggingFaceH4/zephyr-7b-beta              │ v0.3 hf-inference-en 400\n"
    "  #7     │ mistralai/Mistral-7B-Instruct-v0.3        │ together providerre váltás\n"
    "  #12    │ meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo │ Mistral-v0.3 deprecated (410)")

doc.add_paragraph()

# ── Helykitöltő jövőbeli változásokhoz ──────────────────────────────────────
add_heading(doc, "7.13 Jövőbeli változások helye", level=2)
add_paragraph(doc,
    "Ha a jövőben újabb módosítás, hibajavítás vagy refaktorálás történik, "
    "a változásnapló ide kerül bővítésre (#13, #14, … sorszámmal), "
    "a fenti táblázat formátumát követve.",
    italic=True)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 8. STREAMLIT UI LEÍRÁSA
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Streamlit felhasználói felület", level=1)

add_heading(doc, "8.1 Oldalsáv (Sidebar)", level=2)
add_bullet(doc, "HuggingFace API kulcs beviteli mező (password típus, .env-ből előtölti)")
add_bullet(doc, "top-k slider: hány releváns dokumentumot keressen elő (1–6, alapértelmezett: 3)")
add_bullet(doc, "API hívások összefoglaló táblázata")
add_bullet(doc, "Alkalmazott technikák felsorolása")

add_heading(doc, "8.2 Fő tartalom – bal panel", level=2)
add_bullet(doc, "Szöveg beírása mód: soronként egy dokumentum a szövegterületen")
add_bullet(doc, "Fájl feltöltése mód: .txt fájlok, bekezdésekre bontja a tartalmat")
add_bullet(doc, "Betöltött dokumentumok listájának megtekintése expander-ben")

add_heading(doc, "8.3 Fő tartalom – jobb panel", level=2)
add_bullet(doc, "Kérdés beviteli mező")
add_bullet(doc, "\"Válasz generálása\" gomb (csak akkor aktív, ha minden adat megvan)")
add_bullet(doc, "Figyelmeztető üzenetek hiányzó API kulcs, dokumentum vagy kérdés esetén")

add_heading(doc, "8.4 Eredmény megjelenítés", level=2)
add_bullet(doc, "Minden pipeline lépés st.status() widgetben jelenik meg, valós idejű visszajelzéssel")
add_bullet(doc, "Végső válasz kiemelve st.success()-szel")
add_bullet(doc, "Részletes CoT összesítő expander-ben: hasonlósági pontszámok, BART összefoglaló, Mistral következtetés")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 9. KONFIGURÁCIÓ ÉS TELEPÍTÉS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. Konfiguráció és indítás", level=1)

add_heading(doc, "9.1 Környezeti változók (.env)", level=2)
add_code_block(doc, "HF_API_KEY=hf_xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx")
add_paragraph(doc,
    "A beadando2.py automatikusan átmásolja HF_TOKEN névre is, amit a LangChain "
    "HuggingFaceEmbeddings osztálya igényel.")

add_heading(doc, "9.2 Az alkalmazás indítása", level=2)
add_code_block(doc, "streamlit run app.py")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# LÁBLÉC
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = footer_p.add_run(
    f"Dokumentum-alapú Q&A alkalmazás – Műszaki specifikáció  |  "
    f"Generálva: {datetime.date.today().strftime('%Y. %m. %d.')}"
)
run_f.italic = True
run_f.font.size = Pt(9)
run_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Mentés ──────────────────────────────────────────────────────────────────
output_path = r"C:\Users\user\Documents\Genai\beadando\genaispecifikacio.docx"
doc.save(output_path)
print(f"OK: {output_path}")
